"""
DOCX parser — extracts structured text from Word documents.

Uses python-docx.  Headings become '--- Heading Text ---' section
markers so the downstream section-context chunker automatically
groups content under the right heading prefix.

Install:
    pip install python-docx
"""
from __future__ import annotations

import io
import re
from typing import List

from app.services.document_parsers.base import BaseParser, ParsedDocument, ParserError


# Heading styles that become section markers in the output
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3",
    "title", "subtitle",
}

# Paragraph styles that are noise (page numbers, footers, etc.)
_SKIP_STYLES = {
    "header", "footer", "page number", "footnote text",
    "endnote text", "caption",
}


class DocxParser(BaseParser):
    """
    Converts a .docx file into plain text with section markers.

    Paragraphs styled as Heading 1/2/3 become '--- Heading ---' lines,
    which the section-context chunker picks up as section boundaries.

    Parameters
    ----------
    include_tables : bool
        If True, table cells are extracted row by row as pipe-separated text.
    skip_empty : bool
        Skip paragraphs that are entirely whitespace.
    """

    supported_mimetypes  = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )
    supported_extensions = ("docx", "doc")

    def __init__(
        self,
        include_tables: bool = True,
        skip_empty:     bool = True,
    ) -> None:
        self.include_tables = include_tables
        self.skip_empty     = skip_empty

    def parse(self, data: bytes, filename: str = "") -> ParsedDocument:
        try:
            import docx
        except ImportError as exc:
            raise ParserError(
                "python-docx is required: pip install python-docx"
            ) from exc

        try:
            doc = docx.Document(io.BytesIO(data))
        except Exception as exc:
            raise ParserError(f"Cannot open DOCX '{filename}': {exc}") from exc

        # Extract title from core properties or filename
        title = ""
        try:
            title = doc.core_properties.title or ""
        except Exception:
            pass
        if not title:
            title = _infer_title(filename)

        blocks: List[str] = []

        # ── Body paragraphs ───────────────────────────────────────────────
        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ""
            text       = para.text.strip()

            if not text and self.skip_empty:
                continue
            if any(skip in style_name for skip in _SKIP_STYLES):
                continue

            if any(h in style_name for h in _HEADING_STYLES):
                # Convert headings to section markers the chunker understands
                blocks.append(f"\n--- {text} ---\n")
            else:
                blocks.append(text)

        # ── Tables ────────────────────────────────────────────────────────
        if self.include_tables:
            for table in doc.tables:
                table_lines: List[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    # Deduplicate merged cells (python-docx repeats them)
                    seen, unique = set(), []
                    for c in cells:
                        if c not in seen:
                            seen.add(c)
                            unique.append(c)
                    line = " | ".join(unique)
                    if line.replace("|", "").strip():
                        table_lines.append(line)

                if table_lines:
                    blocks.append("\n" + "\n".join(table_lines) + "\n")

        full_text = "\n\n".join(blocks).strip()
        full_text = re.sub(r"\n{3,}", "\n\n", full_text)   # collapse excess blank lines

        if not full_text:
            raise ParserError(f"DOCX '{filename}' contains no extractable text.")

        return ParsedDocument(
            text  = full_text,
            title = title,
            extra = {"source_type": "docx", "filename": filename},
        )


# ── helpers ───────────────────────────────────────────────────────────────────

def _infer_title(filename: str) -> str:
    name = filename.rsplit(".", 1)[0] if "." in filename else filename
    name = re.sub(r"[_\-]+", " ", name)
    return name.strip().title()